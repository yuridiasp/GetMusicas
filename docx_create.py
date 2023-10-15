from docx import Document
from bs4 import BeautifulSoup
import json
import re

with open("json/musicas-html.json", "r", encoding="latin-1") as file:
    string = file.read()
    lista = json.loads(string)

    doc = Document()

    for html in lista:
        para = doc.add_paragraph()
        pattern = "tom:\n"
        soup = BeautifulSoup(html + "\x0C", 'html.parser')
        para.add_run(soup.get_text())

    doc.save("arquivo_word.docx")